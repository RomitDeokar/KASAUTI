"""
Generate the idea-bank deliverable (.docx + .xlsx).

This is a submission-support artifact, not part of KASAUTI itself. It exists
so the ratings can be handed to other reviewers for a second opinion with the
full reasoning attached, rather than just a number.

    python scripts/make_idea_bank.py
"""
from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")

# ---------------------------------------------------------------------------
# THE RATINGS
#
# Scoring model, stated so it can be argued with:
#   Overall = 0.22*Uniqueness + 0.28*Feasibility_1day + 0.22*RubricFit
#           + 0.28*CollisionSafety
#
# Feasibility is weighted at 0.28 -- up from 0.20 in the v4 workbook -- and
# measured against ONE DAY, not eight. That single change is what reorders the
# entire table. Collision Safety keeps its 0.28: shipping something Razorpay
# already sells is still the fastest way to lose.
# ---------------------------------------------------------------------------
IDEAS = [
    dict(
        rank=1, name="KASAUTI (+ consortium layer)", track="05 / 02",
        uniq=9.0, feas=9.5, rubric=9.0, collision=8.0,
        verdict="BUILD THIS — it is already built",
        one_liner=(
            "Executable conformance suite for money-moving AI agents: Indian "
            "regulation compiled into deterministic checkers, an LLM generates "
            "the attacks, pure code issues every verdict. Three scopes: "
            "per-episode, cross-episode, cross-merchant."),
        why=(
            "Feasibility is 9.5 not because the idea is easy but because the "
            "repo exists and passes 371 tests today. On a one-day horizon that "
            "is the only variable that matters. It also absorbs the two ideas "
            "you wanted to merge: RazorGate's enforcement demo is gateway.py, "
            "RIWAAJ's cross-merchant shape is consortium.py — both as "
            "evaluation strategies over ONE predicate set, which is a stronger "
            "story than three separate products.\n\n"
            "The collision argument is the sharpest in the bank: Razorpay's "
            "guardrails blog describes certification (s8) and runtime "
            "validation (s4) as separate systems and never establishes that "
            "they agree. KASAUTI does not rebuild either. It ships the "
            "equivalence proof between them, which is the one artifact an "
            "outsider can construct and Razorpay has not published."),
        risk=(
            "Precision 1.00 on a self-built corpus invites the circularity "
            "attack, and it is a fair one. Mitigated — not eliminated — by "
            "leading with it in the README, by 10 hard negatives, by an "
            "independent LLM adversary, and by NOT_CHECKED.md sections 1-2 "
            "saying it in the author's own words before a judge says it."),
    ),
    dict(
        rank=2, name="RIWAAJ (standalone cross-merchant ring detection)",
        track="02", uniq=8.0, feas=2.0, rubric=8.0, collision=8.0,
        verdict="ABSORBED — do not start fresh",
        one_liner=(
            "Cross-merchant refund/return-abuse RING detection: the fraud "
            "shape a single merchant's data structurally cannot reveal."),
        why=(
            "Genuinely differentiated and correctly ranked co-#1 in the v4 "
            "workbook when eight days remained. With one day it is "
            "unbuildable: entity graph, community detection, calibrated "
            "classifier, held-out topologies and a cost matrix is a week of "
            "work.\n\n"
            "The honest half — the deterministic aggregation whose ground "
            "truth is not a modelling assumption — is now shipped inside "
            "KASAUTI as consortium.py."),
        risk=(
            "The fatal critique is 'did you just detect your own generator?'. "
            "A learned detector trained on synthetic rings cannot answer it. "
            "This is exactly why consortium.py ships the log-derived rules "
            "and explicitly does NOT ship a learned ring detector."),
    ),
    dict(
        rank=3, name="RazorGate (standalone policy firewall)", track="01",
        uniq=7.0, feas=4.0, rubric=7.0, collision=5.0,
        verdict="ABSORBED — and weaker standalone",
        one_liner=(
            "Merchant-side policy firewall between an inbound AI buyer agent "
            "and Razorpay; deterministic rules block what an LLM tries to "
            "authorize, with a live prompt-injection demo."),
        why=(
            "The prompt-injection demo is the most memorable 60 seconds "
            "available in this hackathon, and it is preserved: failure_lab.py "
            "plus the Provenance enum.\n\n"
            "Standalone it is docked hard on collision. Razorpay guardrails "
            "blog s4 IS a platform-level validation layer enforcing merchant "
            "discount ceilings with a full audit trail. Building a second one "
            "is a feature request wearing a project's clothes. Inside KASAUTI "
            "the same code stops being a competing firewall and becomes half "
            "of an equivalence proof — which is not a repositioning trick, it "
            "is a different artifact."),
        risk=(
            "A blocklist-based injection defence loses to paraphrase. "
            "Provenance survives because it never reads the attack — it reads "
            "which source the agent cited as authority."),
    ),
    dict(
        rank=4, name="PARCHI (settlement underpayment forensics)", track="04",
        uniq=7.5, feas=6.0, rubric=8.0, collision=7.0,
        verdict="BEST FALLBACK IF YOU HAD 5+ DAYS",
        one_liner=(
            "Recompute from first principles what the processor owed you "
            "(MDR + GST-on-MDR + TDS + rounding) and flag silent shortfalls — "
            "an adversarial audit of Razorpay itself, not a convenience "
            "matcher."),
        why=(
            "The adversarial framing is a real pivot away from Agentic "
            "Dashboard's convenience reconciliation. Auditing whether the "
            "processor actually paid correctly is a different question from "
            "'do these rows match', and the arithmetic is deterministic, "
            "unit-testable and honest. Strong problem taste."),
        risk=(
            "Needs real settlement-report structure to be credible, and the "
            "demo is visually dry. Not startable on day 13 of 13."),
    ),
    dict(
        rank=5, name="SETU (agent-to-agent dispute adjudication)", track="02/01",
        uniq=8.5, feas=3.0, rubric=7.5, collision=8.0,
        verdict="HIGHEST CEILING, WRONG WEEK",
        one_liner=(
            "AP2-signed Intent→Cart→Payment mandate chains become "
            "machine-readable evidence; a deterministic adjudicator rules "
            "BUYER / SELLER / SPLIT / HUMAN_REQUIRED."),
        why=(
            "Genuinely uncovered — no Agent Studio agent does agent-to-agent "
            "adjudication, and the insight that a signed mandate chain IS the "
            "evidence is the most elegant idea in the whole bank."),
        risk=(
            "Ed25519 mandate plumbing is a full day before any adjudication "
            "logic exists. On a one-day budget you would ship signature "
            "verification and no product."),
    ),
    dict(
        rank=6, name="DASTAK (consent & suppression ledger)", track="03/05",
        uniq=8.0, feas=5.5, rubric=6.5, collision=8.0,
        verdict="PARTIALLY ABSORBED",
        one_liner=(
            "Append-only hash-chained consent ledger: cryptographically "
            "provable 'we were permitted to contact this person', with "
            "permanent opt-out enforcement."),
        why=(
            "Narrow and unglamorous in the way judges respect. The consent and "
            "opt-out rules already live in KASAUTI's checkers, and "
            "SUPPRESSION_BREACH_NETWORK is the multi-merchant version of its "
            "central claim. The v3 workbook's own advice was 'overlaps "
            "KASAUTI — pick one, don't build both'. Still correct."),
        risk=(
            "Standalone, 'consent infrastructure' is one step removed from any "
            "track's literal ask, so it needs an explicit bridging argument."),
    ),
    dict(
        rank=7, name="MEHNAT (hardship-aware recovery)", track="03",
        uniq=7.5, feas=4.0, rubric=6.5, collision=7.5,
        verdict="BEAUTIFUL, UNMEASURABLE IN TIME",
        one_liner=(
            "Detects genuine financial distress and routes AWAY from "
            "collection into restructuring — the agent that chooses not to "
            "collect."),
        why=(
            "Outstanding problem taste, and the exact opposite value "
            "proposition from Subscription Recovery Agent's push-to-collect "
            "design. Would be memorable."),
        risk=(
            "The hardest thing in the bank to measure honestly. Without a "
            "defensible proxy metric it reads as a vibe, and 'we chose not to "
            "collect' with no measurement is indistinguishable from a broken "
            "collector."),
    ),
    dict(
        rank=8, name="Ledger Court (N:M optimization reconciliation)",
        track="04", uniq=7.5, feas=3.0, rubric=8.0, collision=6.5,
        verdict="TIMELINE EATER",
        one_liner=(
            "Treats settlement matching as a min-cost-flow / integer "
            "programming assignment problem rather than row-by-row matching."),
        why=(
            "The most technically impressive recon variant, and the N:M case "
            "genuinely survives the Agentic Dashboard collision."),
        risk=(
            "OR-Tools formulation can consume the entire budget and produce a "
            "solver that is correct on a toy instance and unexplainable on a "
            "real one. PARCHI gets most of the credit for a third of the "
            "risk."),
    ),
    dict(
        rank=9, name="Sentinel (change-point + FDR diagnostic layer)",
        track="03", uniq=8.5, feas=2.0, rubric=8.0, collision=8.0,
        verdict="STATS TRAP",
        one_liner=(
            "CUSUM change-point detection + false-discovery-rate correction "
            "as a diagnostic layer over payment-failure segments."),
        why=(
            "Intellectually the most rigorous idea in the bank, and the "
            "reframe to a diagnostic layer does clear the collision."),
        risk=(
            "Only build this if you could explain Benjamini-Hochberg on a "
            "whiteboard today. Every prior review said the same thing across "
            "three independent rounds. Half-implemented, it collapses into a "
            "generic threshold alert."),
    ),
    dict(
        rank=10, name="LiftLedger (incrementality auditor)", track="03",
        uniq=8.5, feas=2.5, rubric=7.5, collision=8.0,
        verdict="RIGOR TRAP",
        one_liner=(
            "Sits ON TOP of any recovery agent — including Razorpay's own — "
            "and asks whether it causes incremental revenue or just captures "
            "organic recoveries."),
        why=(
            "The reframe is genuinely clever: auditing whether Razorpay's own "
            "shipped agent works is something no Agent Studio agent does, and "
            "the question is one a payments company actually cares about."),
        risk=(
            "Uplift modelling without a real experiment is a causal claim you "
            "cannot defend. Shipping an indefensible causal claim scores worse "
            "than a simple thing done honestly."),
    ),
    dict(
        rank=11, name="RetryRight (e-mandate retry legality)", track="03",
        uniq=6.5, feas=5.5, rubric=7.0, collision=5.0,
        verdict="PARTIAL COLLISION",
        one_liner=(
            "'The law decides IF, the model decides WHEN' — NPCI/RBI e-mandate "
            "retry rules as a unit-tested compliance engine."),
        why=(
            "The regulation-as-code instinct is right, and it is the instinct "
            "KASAUTI is built on. Applied to retry specifically, it lands on "
            "shipped product."),
        risk=(
            "Subscription Recovery Agent already ships 'smarter retry logic'. "
            "A judge will ask why this is not a feature request to that team. "
            "Also requires reading actual circulars — a day you do not have."),
    ),
    dict(
        rank=12, name="RTO Radar", track="02 (+03)",
        uniq=2.0, feas=8.5, rubric=6.5, collision=1.0,
        verdict="DEAD — DOUBLY CONFIRMED",
        one_liner=(
            "Predicts COD Return-to-Origin risk and fires an intervention "
            "ladder sized by expected-value math."),
        why="Was ranked #1 in the original bank, before collision checking.",
        risk=(
            "Magic Checkout's RTO Intelligence AND a dedicated Agent Studio "
            "RTO/COD agent both ship this. Two independent confirmations, two "
            "products, same company, the one judging you."),
    ),
    dict(
        rank=13, name="Chargeback Evidence Responder / CE3.0", track="02",
        uniq=2.0, feas=7.0, rubric=6.0, collision=1.5,
        verdict="DEAD — CONFIRMED",
        one_liner=(
            "Deterministic Visa CE3.0 eligibility engine; LLM drafts the "
            "rebuttal narrative."),
        why=(
            "Ranked 8/10 and 'least crowded lane in the hackathon' in an "
            "earlier round. That round had not checked."),
        risk=(
            "Razorpay's Dispute Responder gathers evidence, scores win "
            "probability and submits. Identical feature set, shipped."),
    ),
    dict(
        rank=14, name="Hinglish voice recovery cluster", track="03",
        uniq=2.0, feas=7.5, rubric=7.0, collision=1.0,
        verdict="DEAD — WORST COMBINATION",
        one_liner=(
            "Failure classifier → channel selector → Hinglish voice/WhatsApp "
            "→ promise-to-pay tracker."),
        why="Proposed independently by 4 of 12 original sources.",
        risk=(
            "Subscription Recovery Agent ships this in Hindi with a voice "
            "vendor pre-wired. Simultaneously the most crowded idea in the "
            "exercise AND a direct collision."),
    ),
    dict(
        rank=15, name="Generic reconciliation agent", track="04",
        uniq=3.0, feas=7.5, rubric=8.0, collision=2.5,
        verdict="DEAD — USE PARCHI INSTEAD",
        one_liner=(
            "Multi-source reconciliation, exact-match first, fuzzy on the "
            "residue, honest exception queue."),
        why="The honest-exception-queue instinct is good; port it to PARCHI.",
        risk=(
            "Agentic Dashboard ships 'upload a bank statement, get instant "
            "reconciliation' as a feature."),
    ),
    dict(
        rank=16, name="Generic conversational checkout bot", track="01",
        uniq=1.0, feas=8.0, rubric=2.0, collision=1.5,
        verdict="DEAD ON ARRIVAL",
        one_liner="Chat-based AI shopping assistant bolted onto checkout.",
        why="The hello-world of this hackathon.",
        risk=(
            "Agentic Payments (UPI Reserve Pay, live) plus the official 40+ "
            "tool MCP server already cover it."),
    ),
]

W_UNIQ, W_FEAS, W_RUBRIC, W_COLL = 0.22, 0.28, 0.22, 0.28


def overall(i: dict) -> float:
    return round(W_UNIQ * i["uniq"] + W_FEAS * i["feas"]
                 + W_RUBRIC * i["rubric"] + W_COLL * i["collision"], 2)


# Rank is DERIVED from the score, never hand-asserted. An earlier draft of
# this file hardcoded a rank field and it immediately disagreed with the
# arithmetic -- RIWAAJ sat at rank 2 while scoring 6.32, below four ideas
# ranked beneath it. A ranking that contradicts its own formula is worse
# than no ranking, because it tells a reviewer the numbers are decoration.
IDEAS.sort(key=lambda i: -overall(i))
for _n, _i in enumerate(IDEAS, 1):
    _i["rank"] = _n


# ---------------------------------------------------------------------------
def _h(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def _p(doc, text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def build_docx(path: str) -> None:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.8)

    _p(doc, "Razorpay AI Buildathon 2026", bold=True, size=24)
    _p(doc, "Idea Bank — Final Ratings, Re-scored Against the Real Deadline",
       bold=True, size=14)
    _p(doc, "Compiled 4 September 2026 · Deadline 5 September 2026 · "
            "Time remaining: ~1 day", italic=True, size=10)

    doc.add_paragraph()
    _p(doc, "THE FINDING THAT REORDERS EVERYTHING", bold=True, size=13)
    _p(doc,
       "Every previous version of this bank scored feasibility against 13 "
       "days, then 8. Today is 4 September. The deadline is 5 September. The "
       "real budget is roughly one day.\n\n"
       "That is not a small adjustment to the weights — it invalidates the "
       "ranking. An 8.5/10 idea you cannot start is worth less than a 7/10 "
       "already running with honest metrics. On a one-day horizon, 'does "
       "working code exist right now' stops being a tiebreaker and becomes "
       "the dominant term.\n\n"
       "The recommendation is therefore not a new idea. It is: ship KASAUTI, "
       "which already exists, passes 371 tests, and — after today's work — "
       "structurally absorbs both RazorGate and RIWAAJ as additional "
       "evaluation scopes over one shared predicate set.")

    doc.add_paragraph()
    _p(doc, "Scoring model", bold=True, size=12)
    _p(doc, "Overall = 0.22·Uniqueness + 0.28·Feasibility(1 day) "
            "+ 0.22·Rubric Fit + 0.28·Collision Safety", italic=True)
    _p(doc,
       "Feasibility rises from 0.20 to 0.28 and is measured against one day, "
       "not eight. Collision Safety holds 0.28 — shipping something Razorpay "
       "already sells is still the fastest way to lose. Collision Safety: "
       "10 = Razorpay ships nothing like it; 1 = they ship it today.", size=9)

    doc.add_page_break()

    # ---- ranking table
    _h(doc, "Full ranking", 1)
    t = doc.add_table(rows=1, cols=8)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ["#", "Idea", "Track", "Uniq", "Feas", "Rubric", "Coll", "OVERALL"]
    for c, txt in zip(t.rows[0].cells, hdr):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(8)
    for i in IDEAS:
        row = t.add_row().cells
        vals = [str(i["rank"]), i["name"], i["track"], f'{i["uniq"]:.1f}',
                f'{i["feas"]:.1f}', f'{i["rubric"]:.1f}',
                f'{i["collision"]:.1f}', f"{overall(i):.2f}"]
        for c, v in zip(row, vals):
            c.text = ""
            r = c.paragraphs[0].add_run(v)
            r.font.size = Pt(8)
            if v == vals[-1]:
                r.bold = True

    doc.add_paragraph()
    _p(doc, "Note the shape of the table: the dead ideas score well on "
            "feasibility and terribly on collision, and the brilliant ideas "
            "score well on uniqueness and terribly on feasibility. Only the "
            "top entry is strong on both, and only because it is already "
            "built.", italic=True, size=9)

    doc.add_page_break()

    # ---- details
    _h(doc, "Idea-by-idea", 1)
    for i in IDEAS:
        _h(doc, f'{i["rank"]}. {i["name"]}  —  {overall(i):.2f}/10', 2)
        _p(doc, f'Track {i["track"]} · {i["verdict"]}', bold=True, size=10)
        _p(doc, i["one_liner"], italic=True, size=10)
        _p(doc, "Why this score", bold=True, size=10)
        _p(doc, i["why"], size=10)
        _p(doc, "The strongest attack on it", bold=True, size=10)
        _p(doc, i["risk"], size=10)
        _p(doc, f'Uniqueness {i["uniq"]}  ·  Feasibility(1d) {i["feas"]}  ·  '
                f'Rubric fit {i["rubric"]}  ·  Collision safety '
                f'{i["collision"]}', size=8, italic=True)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- the merge
    _h(doc, "The merge you asked for", 1)
    _p(doc,
       "You proposed combining KASAUTI, RazorGate and RIWAAJ. That instinct "
       "is right, but the naive version — three products in one repo — would "
       "read as scope sprawl and fail on all four rubric lines at once.\n\n"
       "The version that works treats all three as ONE set of predicates "
       "evaluated at three different scopes. Nothing is duplicated, so "
       "nothing can drift:")
    m = doc.add_table(rows=1, cols=3)
    m.style = "Light Grid Accent 1"
    for c, txt in zip(m.rows[0].cells,
                      ["Original idea", "How it appears in KASAUTI",
                       "Why the merge is stronger than the standalone"]):
        c.text = ""
        r = c.paragraphs[0].add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
    rows = [
        ("KASAUTI",
         "rules/checkers.py — 8 per-episode rules, each citing a regulation",
         "Unchanged: the core conformance suite."),
        ("RazorGate",
         "gateway.py — the same checkers evaluated inline, before each action",
         "Standalone it collides with guardrails blog s4. As a second "
         "evaluation strategy it becomes half of an equivalence proof — a "
         "different artifact, not a repositioning."),
        ("RIWAAJ",
         "consortium.py — 3 network rules over pooled, salted-hash-joined "
         "merchant reports",
         "Standalone it needs a learned ring detector and dies to 'did you "
         "detect your own generator?'. Scoped to log-derived aggregation, the "
         "ground truth is real and the critique does not apply."),
    ]
    for a, b, c in rows:
        cells = m.add_row().cells
        for cell, txt in zip(cells, (a, b, c)):
            cell.text = ""
            r = cell.paragraphs[0].add_run(txt)
            r.font.size = Pt(8)

    doc.add_paragraph()
    _p(doc, "The one-sentence pitch", bold=True, size=12)
    _p(doc,
       "\"Razorpay promises every Agent Studio agent is certified against "
       "dark-pattern and consent rules, and separately promises every action "
       "passes a runtime validation layer — but nothing establishes that the "
       "two agree, and no developer can test either. KASAUTI is the missing "
       "test: one set of deterministic, citation-bearing predicates, "
       "evaluated offline, inline, and across merchants, with an LLM "
       "generating the attacks and never issuing a verdict.\"",
       italic=True, size=11)

    doc.add_page_break()

    # ---- what to do with the remaining day
    _h(doc, "What the remaining time buys", 1)
    _p(doc, "Build is done. Everything below is submission surface, which is "
            "where the marginal point actually is.", italic=True, size=10)
    for txt in [
        "1. Record the 5-minute video. 30s problem → 90s live demo (make "
        "demo, then make consortium) → 60s the AI-judgment split → 60s "
        "FAILURES.md #9, the phantom-join bug → 30s NOT_CHECKED.md. Ending on "
        "limitations reads as senior, not weak.",
        "2. Fresh-clone test. git clone into a clean directory, pip install "
        "-r requirements.txt, make demo. If that fails on the judge's machine "
        "nothing else matters.",
        "3. Write 'what broke' from FAILURES.md #9. They read this answer "
        "first. It is real, it is specific, it was found by tooling you built "
        "to attack your own work, and it is pinned by 20 regression tests. Do "
        "not embellish it.",
        "4. Submit on the 4th, not the 5th.",
    ]:
        _p(doc, txt, size=10)

    doc.add_paragraph()
    _p(doc, "Open-source and free tools actually used", bold=True, size=12)
    for txt in [
        "Hypothesis — property-based testing; found real bugs in the rules "
        "(FAILURES.md #3), not decoration.",
        "pytest — 371 tests. The only two dependencies in requirements.txt.",
        "Gemini free tier — the LLM adversary only. Offline seeded mutation "
        "is the default, so the demo reproduces with no key and no network.",
        "Python stdlib (hashlib, ast, dis) — the purity guard and the "
        "consortium join. Zero runtime dependencies is itself a build-quality "
        "signal.",
        "Cited, not vendored: razorpay/razorpay-mcp-server, "
        "google-agentic-commerce/AP2, ethz-spylab/agentdojo, "
        "invariantlabs-ai/invariant (the comparison baseline).",
    ]:
        _p(doc, "· " + txt, size=9)

    doc.add_paragraph()
    _p(doc, "For reviewers being asked for a second opinion", bold=True,
       size=12)
    _p(doc,
       "The four questions worth attacking, in the order they would hurt "
       "most:\n\n"
       "1. Is precision 1.00 on a self-built corpus worth anything? (My "
       "answer: only as self-consistency plus a fair baseline comparison — "
       "NOT_CHECKED.md §1-2 says so before a judge does.)\n"
       "2. Does the certification-vs-enforcement gap actually exist, or is it "
       "inferred from a blog post? (Inferred, and stated as such in "
       "NOT_CHECKED.md §4.)\n"
       "3. Is a salted hash meaningfully privacy-preserving at consortium "
       "scale? (No — NOT_CHECKED.md §6 enumerates exactly how it fails.)\n"
       "4. Is 'we chose not to build the ML component' rigour or an excuse? "
       "(The strongest genuine objection. My answer is that a detector scored "
       "against rings I invented would be worth less than the arithmetic that "
       "is provably right — but a reviewer could reasonably disagree.)",
       size=10)

    doc.save(path)


def build_xlsx(path: str) -> None:
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError:
        print("openpyxl not installed; skipping xlsx")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Idea Rankings"
    headers = ["Rank", "Idea", "Track", "One-Liner", "Uniqueness",
               "Feasibility (1 day)", "Rubric Fit", "Collision Safety",
               "OVERALL", "Verdict", "Why this score",
               "Strongest attack on it"]
    ws.append(headers)
    hf = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="1F3864")
    for c in ws[1]:
        c.font = hf
        c.fill = fill
        c.alignment = Alignment(vertical="center", wrap_text=True)
    for i in IDEAS:
        ws.append([i["rank"], i["name"], i["track"], i["one_liner"],
                   i["uniq"], i["feas"], i["rubric"], i["collision"],
                   overall(i), i["verdict"], i["why"], i["risk"]])
    widths = [6, 34, 10, 52, 11, 13, 10, 13, 10, 30, 70, 60]
    for col, w in zip("ABCDEFGHIJKL", widths):
        ws.column_dimensions[col].width = w
    for row in ws.iter_rows(min_row=2):
        for c in row:
            c.alignment = Alignment(vertical="top", wrap_text=True)

    ws2 = wb.create_sheet("Read Me")
    for line in [
        ["Razorpay AI Buildathon 2026 - Idea Bank, re-scored 4 Sept 2026"],
        [""],
        ["THE FINDING: previous versions scored feasibility against 13 days, "
         "then 8. Today is 4 Sept; the deadline is 5 Sept. ~1 day remains."],
        ["That invalidates the old ranking. An 8.5 you cannot start loses to "
         "a 7 already running with honest metrics."],
        [""],
        ["Overall = 0.22*Uniqueness + 0.28*Feasibility(1d) + 0.22*RubricFit "
         "+ 0.28*CollisionSafety"],
        ["Collision Safety: 10 = Razorpay ships nothing like it, "
         "1 = they ship it today."],
        [""],
        ["RECOMMENDATION: ship KASAUTI. It exists, passes 371 tests, and now "
         "absorbs RazorGate (gateway.py) and RIWAAJ (consortium.py) as "
         "additional evaluation scopes over ONE predicate set."],
        [""],
        ["Repo: https://github.com/RomitDeokar/razorpay"],
    ]:
        ws2.append(line)
    ws2.column_dimensions["A"].width = 120
    for row in ws2.iter_rows():
        for c in row:
            c.alignment = Alignment(wrap_text=True, vertical="top")
    ws2["A1"].font = Font(bold=True, size=14)

    wb.save(path)


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    d = os.path.join(OUT_DIR, "razorpay-buildathon-idea-bank-FINAL-v5.docx")
    x = os.path.join(OUT_DIR, "razorpay-buildathon-idea-bank-FINAL-v5.xlsx")
    build_docx(d)
    build_xlsx(x)
    print("wrote", d)
    print("wrote", x)
    print()
    print(f"{'#':>3}  {'IDEA':<44} {'OVERALL':>8}")
    for i in sorted(IDEAS, key=lambda z: -overall(z)):
        print(f'{i["rank"]:>3}  {i["name"][:44]:<44} {overall(i):>8.2f}')
